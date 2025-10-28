#
# generator workflow class used by Django app and command line
#
import sys
import logging
from logging import Logger
import json
import time
from datetime import datetime

from pydantic import BaseModel, Field
import pydantic_ai
from pydantic_ai import Agent
from pydantic_ai.models.openai import OpenAIModel
from pydantic_ai.providers.openai import OpenAIProvider
from pydantic_ai.usage import Usage

import chromadb
from chromadb import Collection
from chromadb.config import DEFAULT_TENANT, DEFAULT_DATABASE, Settings
from chromadb.utils.embedding_functions import OllamaEmbeddingFunction

from openai import OpenAI

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# local
from common import OneRecord, OneDesc, AllDesc, OneQueryResult, OneResultList, ConfigSingleton, DebugUtils, ReportIssue, AllReportIssues, OpenFile, OneEmployer, AllEmployers
from workflowbase import WorkflowBase 

class GeneratorWorkflow(WorkflowBase):

    def __init__(self, context : dict, logger : Logger):
        """
        Args:
            context (dict)
            logger (Logger) - can originate in CLI or Django app
        """
        super().__init__(context, logger)


    def threadWorker(self):
        """
        Workflow to read, parse, generate
        
        Args:
            None

        Returns:
            None
        """

        #-----------------stage configure

        start = time.time()
        totalStart = start

        self._context["stage"] = "configure"
        self.workerSnapshot(None)

        # read ad text if CLI, else it is already in context["adtext"]
        if "adtext" in self._context:
            jobAdRecord = OneRecord(
                id = "", 
                name=self._context['adFileName'], 
                description=self._context["adtext"]
            )
            msg = f"Got job descriptions from HTML form"
            self.workerSnapshot(msg)
        else:
            boolResult, contentJDOrError = OpenFile.open(filePath = self._context['adFileName'], readContent = True)
            if not boolResult:
                self.workerError(contentJDOrError)
                return
            jobAdRecord = OneRecord(
                id = "", 
                name=self._context['adFileName'], 
                description=contentJDOrError
            )
            msg = f"Read job descriptions from file {self._context['adFileName']}"
            self.workerSnapshot(msg)

        try:
            chromaClient = chromadb.PersistentClient(
                path=ConfigSingleton().getAbsPath("rag_datapath"),
                settings=Settings(anonymized_telemetry=False),
                tenant=DEFAULT_TENANT,
                database=DEFAULT_DATABASE,
            )
        except Exception as e:
            msg = f"Error: OpenAI API exception: {e}"
            self.workerError(msg)
            return
        
        ef = OllamaEmbeddingFunction(
            model_name=self._config["rag_embed_llm"],
            url=self._config["rag_embed_url"]    
        )

        collectionName = "actreal"
        try:
            chromaActivity = chromaClient.get_collection(
                name=collectionName,
                embedding_function=ef
            )
        except Exception as e:
            msg = f"Error: collection ACTIVITY exception: {e}"
            self.workerError(msg)
            return

        collectionName = "scenario"
        try:
            chromaScenario = chromaClient.get_collection(
                name=collectionName,
                embedding_function=ef
            )
        except Exception as e:
            msg = f"Error: collection SCENARIO exception: {e}"
            self.workerError(msg)
            return

        end = time.time()
        msg = f"Opened vector collections ACTIVITY with {chromaActivity.count()} documents, SCENARIO with {chromaScenario.count()} documents. {(end-start):9.4f} seconds"
        self.workerSnapshot(msg)

        #----------------stage summary

        start = time.time()

        self._context["stage"] = "summary"
        self.workerSnapshot(None)

        execSummary, usageStats = self.extractExecSection(jobAdRecord)
        if not execSummary:
            return

        self._context['jobtitle'] = execSummary.title
        self._context['execsummary'] = execSummary.description
        if usageStats:
            self._context["llmrequests"] = usageStats.requests
            self._context["llmrequesttokens"] = usageStats.request_tokens
            self._context["llmresponsetokens"] = usageStats.response_tokens

        end = time.time()

        if usageStats:
            msg = f"Extracted executive summary from job description. {(end-start):9.4f} seconds. {usageStats.request_tokens} request tokens. {usageStats.response_tokens} response tokens."
        else:
            msg = f"Extracted executive summary from job description. {(end-start):9.4f} seconds."
        self.workerSnapshot(msg)

        #----------------stage extract

        start = time.time()

        self._context["stage"] = "extract"
        self.workerSnapshot(None)

        allDescriptions = AllDesc(
            ad_name = jobAdRecord.name,
            exec_section = execSummary,
            project_list = [])

        oneResultList, usageStats = self.extractInfoFromJobAd(jobAdRecord)

        if not oneResultList:
            msg = f"Internal error on extracting activities from job description"
            self.workerError(msg)
            return

        self._context['extracted'] = oneResultList.results_list
        if usageStats:
            self._context["llmrequests"] += usageStats.requests
            self._context["llmrequesttokens"] += usageStats.request_tokens
            self._context["llmresponsetokens"] += usageStats.response_tokens

        end = time.time()

        if usageStats:
            msg = f"Extracted {len(oneResultList.results_list)} activities from job description. {(end-start):9.4f} seconds. {usageStats.request_tokens} request tokens. {usageStats.response_tokens} response tokens."
        else:
            msg = f"Extracted {len(oneResultList.results_list)} activities from job description. {(end-start):9.4f} seconds."
        self.workerSnapshot(msg)

        #--------------stage mapping

        start = time.time()

        self._context["stage"] = "mapping"
        self.workerSnapshot(None)

        # ChromaDB calls do not account for LLM usage
        oneResultList = self.mapToActivity(oneResultList, chromaActivity)

        self._context['mapped'] = oneResultList.results_list

        end = time.time()

        msg = f"Mapped {len(oneResultList.results_list)} activities to vector database. {(end-start):9.4f} seconds"
        self.workerSnapshot(msg)

        #----------------stage projects

        startAllProjects = time.time()

        self._context["stage"] = "projects"
        self.workerSnapshot(None)

        self._context['projects'] = []
        prjCount = 0
        for chromaQuery in oneResultList.results_list:

            if chromaQuery[:4] == "--- ":
                self._logger.info(f"!!!!----!!!!!---skipping '{chromaQuery}'")
                continue
            start = time.time()
            oneDesc, usageStats = self.makeProject(chromaQuery, chromaScenario)
            if oneDesc:
                prjCount += 1
                allDescriptions.project_list.append(oneDesc)
                self._context['projects'].append(oneDesc.description)

                if usageStats:
                    self._context["llmrequests"] += usageStats.requests
                    self._context["llmrequesttokens"] += usageStats.request_tokens
                    self._context["llmresponsetokens"] += usageStats.response_tokens

                end = time.time()

                msg = f"Project # {prjCount}: {oneDesc.title}. {(end-start):9.4f} seconds. {usageStats.request_tokens} request tokens. {usageStats.response_tokens} response tokens."
                self.workerSnapshot(msg)

        endAllProjects = time.time()

        msg = f"Created {len(self._context['projects'])} projects. {(endAllProjects-startAllProjects):9.4f} seconds"
        self.workerSnapshot(msg)


        #--------------stage completed

        totalEnd = time.time()

        self._context["stage"] = "completed"
        msg = f"Processing completed. Total time {(totalEnd-totalStart):9.4f} seconds. {self._context["llmrequests"]} LLM requests. {self._context["llmrequesttokens"]} request tokens. {self._context["llmresponsetokens"]} response tokens."
        self.workerSnapshot(msg)


    def extractExecSectionOllama(self, jobInfo : OneRecord)  -> tuple[OneDesc, Usage] :
        """
        Extract summary from job add text
        
        Args:
            jobInfo (OneRecord) - description of job add

        Returns:
            OneDesc record and Usage or None
        """

        systemPromptPhase1 = f"""
            You are an expert in cyber security, information technology and software development 
            You will be supplied text of job advertisement.
            Your job is to extract information from the text that matches user's request.
            Here is the JSON schema for the OneDesc model you must use as context for what information is expected:
            {json.dumps(OneDesc.model_json_schema(), indent=2)}
            """

        ollModel = OpenAIModel(model_name=self._config["main_llm_name"], 
                            provider=OpenAIProvider(base_url=self._config["llm_base_url"]))
        agent = Agent(ollModel, 
                    output_type=OneDesc,                  
                    system_prompt = systemPromptPhase1,
                    retries=10,
                    output_retries=10)

        promptPhase1 = f"""Extract the required role suitable for CV from the text below.
        Make required role the title.
        Fill description with generic description of the required role in past tense suitable for CV.
        Do not add formatting.
        Output only the result.
        \n
        {jobInfo.description}
                """

        try:
            result = agent.run_sync(promptPhase1)
            runUsage = result.usage()
            oneDesc = OneDesc.model_validate_json(result.output.model_dump_json())
            for attr in oneDesc.__dict__:
                oneDesc.__dict__[attr] = oneDesc.__dict__[attr].replace("\n", " ")
                oneDesc.__dict__[attr] = oneDesc.__dict__[attr].encode("ascii", "ignore").decode("ascii")
#            DebugUtils.dumpPydanticObject(oneDesc, "Executive Summary")
        except pydantic_ai.exceptions.UnexpectedModelBehavior as e:
            msg = f"extractExecSection: Skipping due to exception: {e}"
            self.workerError(msg)
            return None, None
        return oneDesc, runUsage


    def extractExecSectionGemini(self, jobInfo : OneRecord)  -> tuple[OneDesc, Usage] :
        """
        Use Google Gemini AI Agent to extracts summary from job add text
        
        Args:
            jobInfo (OneRecord) - description of job add

        Returns:
            OneDesc record and Usage or None
        """

        systemPrompt = """
            You are an expert in cyber security, information technology and software development 
            You will be supplied text of job advertisement.
            Your job is to extract information from the text that matches user's request.
            Here is the JSON schema for the OneDesc model you must use as context for what information is expected:
            {json.dumps(OneDesc.model_json_schema(), indent=2)}
        """

        userPrompt = f"""
        Extract the required role suitable for CV from the text below.
        Make required role the title.
        Fill description with generic description of the required role in past tense suitable for CV.
        Do not add formatting.
        Output only the result.
        \n
        {jobInfo.description}
        """

        api_key = self._config["gemini_key"]

        openAIClient = OpenAI(
            api_key=api_key,
            base_url="https://generativelanguage.googleapis.com/v1beta/openai/"
        )

        completion = openAIClient.beta.chat.completions.parse(
            model="gemini-2.0-flash",
            messages=[
                {"role": "system", "content": systemPrompt},
                {"role": "user", "content": userPrompt},
            ],
            response_format=OneDesc,
        )

        oneDesc = completion.choices[0].message.parsed
        for attr in oneDesc.__dict__:
            oneDesc.__dict__[attr] = oneDesc.__dict__[attr].replace("\n", " ")
            oneDesc.__dict__[attr] = oneDesc.__dict__[attr].encode("ascii", "ignore").decode("ascii")

        # map Google usage to Pydantic usage            
        usage = Usage()
        usage.requests = 1
        usage.request_tokens = completion.usage.prompt_tokens
        usage.response_tokens = completion.usage.completion_tokens

        return oneDesc, usage


    def extractExecSection(self, jobInfo : OneRecord)  -> tuple[OneDesc, Usage] :
        if self._context["llmProvider"] == "Ollama":
            oneDesc, usageStats = self.extractExecSectionOllama(jobInfo)
        if self._context["llmProvider"] == "Gemini":
            oneDesc, usageStats = self.extractExecSectionGemini(jobInfo)
        return oneDesc, usageStats


    def extractInfoFromJobAdOllama(self, jobInfo : OneRecord)  -> tuple[OneResultList, Usage] :
        """
        Extract information from job add text
        
        Args:
            jobInfo (OneRecord) - description of job add

        Returns:
            OneResultList and Usage or None
        """

        systemPromptPhase1 = f"""
            You are an expert in cyber security, information technology and software development 
            You will be supplied text of job advertisement.
            Your job is to extract information from the text that matches user's request.
            """

        ollModel = OpenAIModel(model_name=self._config["main_llm_name"], 
                            provider=OpenAIProvider(base_url=self._config["llm_base_url"]))
        agent = Agent(ollModel, 
                    system_prompt = systemPromptPhase1,
                    retries=5,
                    output_retries=5)

        promptPhase1 = f"""Extract the list of 
        activities, technologies, methodologies, software services, and software products from the text below.
        Combine all items in common list of strings.
        Do not separate by category.
        Avoid single word items.
        Output only lower-case characters.
        Output only the result.
        \n
        {jobInfo.description}
        """

        try:
            result = agent.run_sync(promptPhase1)
            runUsage = result.usage()

        except pydantic_ai.exceptions.UnexpectedModelBehavior as e:
            msg = f"extractInfoFromJobAd: Skipping due to exception: {e}"
            self.workerError(msg)
            return None, None

        phase2Input = result.output

        systemPromptPhase2 = f"""
            You are an expert in JSON processing.
            Input prompt contains list in JSON format. 

            Here is the JSON schema for the OneResultList model you must use as context for what information is expected:
            {json.dumps(OneResultList.model_json_schema(), indent=2)}
            """
        promptPhase2 = f"""
                    {phase2Input}
                """

        agentPhase2 = Agent(ollModel, 
                    output_type=OneResultList,
                    system_prompt = systemPromptPhase2,
                    retries=5,
                    output_retries=5)
        try:
            result = agentPhase2.run_sync(promptPhase2)
            oneResultList = OneResultList.model_validate_json(result.output.model_dump_json())
            runUsage += result.usage()
#            DebugUtils.dumpPydanticObject(oneResultList, "list from job ad")
        except pydantic_ai.exceptions.UnexpectedModelBehavior as e:
            msg = f"extractInfoFromJobAd: Skipping due to exception: {e}"
            self.workerError(msg)
            return None, None

        values = []
        for value in oneResultList.results_list:
            value = value.replace("\n", " ")
            value = value.encode("ascii", "ignore").decode("ascii")
            values.append(value)
        oneResultList.results_list = values

        return oneResultList, runUsage


    def extractInfoFromJobAdGemini(self, jobInfo : OneRecord)  -> tuple[OneResultList, Usage] :
        """
        Extract information from job add text using Google Gemini API
        
        Args:
            jobInfo (OneRecord) - description of job add

        Returns:
            OneResultList and Usage or None
        """

        systemPrompt = f"""
            You are an expert in cyber security, information technology and software development 
            You will be supplied text of job advertisement.
            Your job is to extract information from the text that matches user's request.
            Here is the JSON schema for the OneResultList model you must use as context for what information is expected:
            \n
            {json.dumps(OneResultList.model_json_schema(), indent=2)}
            """

        userPrompt = f"""Extract the list of 
        activities, technologies, methodologies, software services, and software products from the text below.
        Combine all items in common list of strings.
        Do not separate by category.
        Avoid single word items.
        Output only lower-case characters.
        Output only the result.
        \n
        {jobInfo.description}
        """
        api_key = self._config["gemini_key"]

        openAIClient = OpenAI(
            api_key=api_key,
            base_url="https://generativelanguage.googleapis.com/v1beta/openai/"
        )

        completion = openAIClient.beta.chat.completions.parse(
            model="gemini-2.0-flash",
            messages=[
                {"role": "system", "content": systemPrompt},
                {"role": "user", "content": userPrompt},
            ],
            response_format=OneResultList,
        )

        oneResultList = completion.choices[0].message.parsed
        values = []
        for value in oneResultList.results_list:
            value = value.replace("\n", " ")
            value = value.encode("ascii", "ignore").decode("ascii")
            values.append(value)
        oneResultList.results_list = values

        # map Google usage to Pydantic usage            
        usage = Usage()
        usage.requests = 1
        usage.request_tokens = completion.usage.prompt_tokens
        usage.response_tokens = completion.usage.completion_tokens

        return oneResultList, usage


    def extractInfoFromJobAd(self, jobInfo : OneRecord)  -> tuple[OneResultList, Usage] :
        """
        Extract information from job add text
        
        Args:
            jobInfo (OneRecord) - description of job add

        Returns:
            OneResultList and Usage or None
        """
        if self._context["llmProvider"] == "Ollama":
            oneResultList, usageStats = self.extractInfoFromJobAdOllama(jobInfo)
        if self._context["llmProvider"] == "Gemini":
            oneResultList, usageStats = self.extractInfoFromJobAdGemini(jobInfo)
        return oneResultList, usageStats


    def mapToActivity(self, oneResultList : OneResultList, chromaDBCollection : Collection) -> OneResultList :
        """
        Map information to activity table. Interaction with local ChromaDB does not use tokens.
        
        Args:
            oneResultList (OneResultList) - results extracted from job ad
            chromaDBCollection (Collection) - activity table

        Returns:
            OneResultList or None
        """
        itemSet = set()
        for itm in oneResultList.results_list:
            listNew = self.getChromaDBMatchActivity(chromaDBCollection, itm)
            for itemNew in listNew.results_list:
                itemSet.add(itemNew)
        oneResultList = OneResultList(results_list = list(itemSet))
#        DebugUtils.dumpPydanticObject(oneResultList, "Mapped list")
        return oneResultList


    def getChromaDBMatchActivity(self, chromaDBCollection : Collection, queryString : str) -> OneResultList :
        """
        Query ChromaDB collection and select vectors within cut-off distance
        
        Args:
            chromaDBCollection (Collection) - vector table
            queryString (str) - query

        Returns:
            OneResultList or None
        """

        totals = set()

        queryResult = chromaDBCollection.query(query_texts=[queryString], n_results=1)
        cutDist = self._config["rag_distmatch"]
        resultIdx = -1
        for distFloat in queryResult["distances"][0] :
            resultIdx += 1
            docText = ""
            if (queryResult["documents"]) :
                docText = queryResult["documents"][0][resultIdx]

            if (distFloat > cutDist) :
                break

            totals.add(docText)

        if not len(totals) :
            self._logger.info(f"Query {queryString} did not get matches less than {self._config['rag_distmatch']}")
            return OneResultList(results_list = [])

        return OneResultList(results_list=list(totals))



    def makeProjectOllama(self, chromaQuery : str, chromaScenario : Collection)  -> tuple[OneDesc, Usage] :
        """
        Make a project from information in scenario table using Ollama host
        
        Args:
            chromaQuery (str) - query to match in scenario table
            chromaScenario (Collection) - scenario table

        Returns:
            OneDesc and Usage or None
        """

        queryResult = chromaScenario.query(query_texts=[chromaQuery], n_results=3)

        idx = -1
        numberChosen = 0
        distList = []
        combinedDoco = ""
        for distFloat in queryResult["distances"][0] :
            idx += 1
            if (distFloat > self._config['rag_scenario']) : 
                break

            distList.append(distFloat)
            docText = ""
            if (queryResult["documents"]) :
                docText = queryResult["documents"][0][idx]
            metaInf = ""
            if (queryResult["metadatas"]) :
                metaInf = queryResult["metadatas"][0][idx]["docName"]
            combinedDoco = combinedDoco + "\n" + docText
            numberChosen = numberChosen + 1

        if not numberChosen :
            msg = f"ERROR: cannot find ChromaDB records under distance {self._config['rag_scenario']}"
            self.workerError(msg)
            return None, None
        
        systemPrompt = f"""
            You are an expert technical writer. 
            Create title and description of the project based on information supplied. 
            Create a full paragraph for description.
            Do not format text. Remove line feeds and carriage returns.
            Output only the result.
            Here is the JSON schema for the OneDesc model you must use as context for what information is expected:
            {json.dumps(OneDesc.model_json_schema(), indent=2)}
            """
        prompt = f"{combinedDoco}"
        # self._logger.info(f"-------prompt---------\n{prompt}\n-------------\n")

        ollModel = OpenAIModel(model_name=self._config["main_llm_name"], 
                            provider=OpenAIProvider(base_url=self._config["llm_base_url"]))
        
        agent = Agent(ollModel,
                    output_type=OneDesc,
                    system_prompt = systemPrompt,
                    retries=5,
                    output_retries=5)
        try:
            result = agent.run_sync(prompt)
            oneDesc = OneDesc.model_validate_json(result.output.model_dump_json())
            oneDesc.description = oneDesc.description.replace("\n", "")
            runUsage = result.usage()
#            DebugUtils.dumpPydanticObject(oneDesc, "Project Description")
        except pydantic_ai.exceptions.UnexpectedModelBehavior:
            msg = f"Exception: pydantic_ai.exceptions.UnexpectedModelBehavior"
            self.workerError(msg)
            return None, None
        
        for attr in oneDesc.__dict__:
            oneDesc.__dict__[attr] = oneDesc.__dict__[attr].replace("\n", " ")
            oneDesc.__dict__[attr] = oneDesc.__dict__[attr].encode("ascii", "ignore").decode("ascii")

        return oneDesc, runUsage


    def makeProjectGemini(self, chromaQuery : str, chromaScenario : Collection)  -> tuple[OneDesc, Usage] :
        """
        Make a project from information in scenario table using Google Gemini API
        
        Args:
            chromaQuery (str) - query to match in scenario table
            chromaScenario (Collection) - scenario table

        Returns:
            OneDesc and Usage or None
        """

        queryResult = chromaScenario.query(query_texts=[chromaQuery], n_results=3)

        idx = -1
        numberChosen = 0
        distList = []
        combinedDoco = ""
        for distFloat in queryResult["distances"][0] :
            idx += 1
            if (distFloat > self._config['rag_scenario']) : 
                break

            distList.append(distFloat)
            docText = ""
            if (queryResult["documents"]) :
                docText = queryResult["documents"][0][idx]
            metaInf = ""
            if (queryResult["metadatas"]) :
                metaInf = queryResult["metadatas"][0][idx]["docName"]
            combinedDoco = combinedDoco + "\n" + docText
            numberChosen = numberChosen + 1

        if not numberChosen :
            msg = f"ERROR: cannot find ChromaDB records under distance {self._config['rag_scenario']}"
            self.workerError(msg)
            return None, None
        
        systemPrompt = f"""
            You are an expert technical writer. 
            Create title and description of the project based on information supplied. 
            Create a full paragraph for description.
            Do not format text. Remove line feeds and carriage returns.
            Output only the result.
            Here is the JSON schema for the OneDesc model you must use as context for what information is expected:
            {json.dumps(OneDesc.model_json_schema(), indent=2)}
            """
        userPrompt = f"{combinedDoco}"

        api_key = self._config["gemini_key"]

        openAIClient = OpenAI(
            api_key=api_key,
            base_url="https://generativelanguage.googleapis.com/v1beta/openai/"
        )

        completion = openAIClient.beta.chat.completions.parse(
            model="gemini-2.0-flash",
            messages=[
                {"role": "system", "content": systemPrompt},
                {"role": "user", "content": userPrompt},
            ],
            response_format=OneDesc,
        )

        oneDesc = completion.choices[0].message.parsed
        for attr in oneDesc.__dict__:
            oneDesc.__dict__[attr] = oneDesc.__dict__[attr].replace("\n", " ")
            oneDesc.__dict__[attr] = oneDesc.__dict__[attr].encode("ascii", "ignore").decode("ascii")

        # map Google usage to Pydantic usage            
        usage = Usage()
        usage.requests = 1
        usage.request_tokens = completion.usage.prompt_tokens
        usage.response_tokens = completion.usage.completion_tokens

        return oneDesc, usage


    def makeProject(self, chromaQuery : str, chromaScenario : Collection)  -> tuple[OneDesc, Usage] :
        """
        Make a project from information in scenario table
        
        Args:
            chromaQuery (str) - query to match in scenario table
            chromaScenario (Collection) - scenario table

        Returns:
            OneDesc and Usage or None
        """
        if self._context["llmProvider"] == "Ollama":
            oneResultList, usageStats = self.makeProjectOllama(chromaQuery, chromaScenario)
        if self._context["llmProvider"] == "Gemini":
            # fit Rate Per Minute quota for free account
            time.sleep(10)
            oneResultList, usageStats = self.makeProjectGemini(chromaQuery, chromaScenario)
        return oneResultList, usageStats



    def makeWordDoc(self, allDesc : AllDesc) :
        """
        Make a Word document with resume
        
        Args:
            allDesc (AllDesc) - complete record

        Returns:
            None
        """

        allEmployers = AllEmployers(employer_list = [])
        allEmployers.employer_list.append(OneEmployer(
            name = "Darlowie Security Consulting",
            blurb = "Principal Consultant",
            start = datetime(2024, 11, 1),
            end = datetime.now(),
            numJobs = 3
        ))
        allEmployers.employer_list.append(OneEmployer(
            name = "NCC Group APAC",
            blurb = "Principal Consultant",
            start = datetime(2014, 12, 1),
            end = datetime(2024, 10, 30),
            numJobs = 15
        ))
        allEmployers.employer_list.append(OneEmployer(
            name = "NCC Group USA",
            blurb = "Senior Consultant",
            start = datetime(2013, 5, 15),
            end = datetime(2014, 10, 30),
            numJobs = 10
        ))

        # Create a new Document
        doc = Document()

        title = doc.add_heading(level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titleRun = title.add_run('Anton Pavlov')
        font = titleRun.font
        font.Name = 'Cambria'
        font.size = Pt(20)

        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run('Melbourne, Victoria  Ph: 0428065430  Email:XXXX@YYYY.com')

        exec_paragraph = doc.add_paragraph()
        exec_paragraph.add_run(allDesc.exec_section.title).bold = True
        exec_paragraph.add_run('\n')
        exec_paragraph.add_run(allDesc.exec_section.description)

        expHeading = doc.add_heading(level=2)
        expHeading.add_run('Experience')

        idxCurrentJob = 0
        idxGlobalJob = 0
        for emp in allEmployers.employer_list:
            date_name_para = doc.add_paragraph()
            date_name_para.add_run(emp.start.strftime('%b %Y') + '-' +  emp.end.strftime('%b %Y'))
            date_name_para.add_run(' ')
            date_name_para.add_run(emp.name).bold = True
            blurb_para = doc.add_paragraph()
            blurb_para.add_run(emp.blurb)
            for oneDesc in allDesc.project_list[idxGlobalJob:] :
                idxCurrentJob += 1
                if (idxCurrentJob > emp.numJobs) :
                    idxCurrentJob = 0
                    break
                list_paragraph = doc.add_paragraph(style='List Bullet')
                list_paragraph.add_run(oneDesc.title).bold = True
                list_paragraph.add_run('\n')
                list_paragraph.add_run(oneDesc.description)
                list_paragraph.add_run('\n')
                idxGlobalJob += 1

        # Save the document
        doc.save(self._context['wordFileName'])

