#
#
#
#

import sys
import tomli
import json
from datetime import datetime


import chromadb
from chromadb import Collection
from chromadb.config import DEFAULT_TENANT, DEFAULT_DATABASE, Settings
from chromadb.utils.embedding_functions import OllamaEmbeddingFunction
from chromadb import QueryResult

import pydantic_ai
from pydantic import BaseModel, Field
from pydantic_ai import Agent
from pydantic_ai.models.openai import OpenAIModel
from pydantic_ai.providers.openai import OpenAIProvider

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# local
from common import OneRecord, AllRecords, OneQueryResult, AllQueryResults, ConfigSingleton, OpenFile
from common import DebugUtils, OneDesc, AllDesc, OneResultList, OneEmployer, AllEmployers

#---------------------------------------------------

def main():

    if len(sys.argv) < 2:
        print(f"Usage:\n\t{sys.argv[0]} FILENAME\nExample: {sys.argv[0]} TEMP001.txt")
        return
    jobAdName = sys.argv[1]

    configName = "default.toml"
    try:
        with open(configName, mode="rb") as fp:
            ConfigSingleton().conf = tomli.load(fp)
    except Exception as e:
        print(f"***ERROR: Cannot open config file {configName}, exception {e}")
        return

    chromaClient = chromadb.PersistentClient(
        path=ConfigSingleton().conf["rag_datapath"],
        settings=Settings(),
        tenant=DEFAULT_TENANT,
        database=DEFAULT_DATABASE,
    )

    chromaActivity = getChromaCollection(chromaClient, "actreal")
    if not chromaActivity:
        return
    
    chromaScenario = getChromaCollection(chromaClient, "scenario")
    if not chromaScenario:
        return

    # read in job ads
    boolResult, contentJDOrError = OpenFile.open(filePath = jobAdName, readContent = True)
    if not boolResult:
        print(contentJDOrError)
        return
    jobAdRecord = OneRecord(
        id = "", 
        name=str(jobAdName), 
        description=contentJDOrError
    )
    print(f"Read in job descriptions from {jobAdName}")

    execSection = extractExecSection(jobAdRecord)

    allDescriptions = AllDesc(
        ad_name = jobAdRecord.name,
        exec_section = execSection,
        project_list = [])

    oneResultList = extractInfoFromJobAd(jobAdRecord)
    if not oneResultList:
        return
    oneResultList = mapToActivity(oneResultList, chromaActivity)
#    oneResultList = mapToActivityFake(oneResultList, chromaActivity)
    
    for chromaQuery in oneResultList.results_list:
        oneDesc = makeProject(chromaQuery, chromaScenario)
        if oneDesc:
            allDescriptions.project_list.append(oneDesc)
#       if DebugUtils.pressKey("Press c to move to next project:"):
#          continue

    makeWordDoc(allDescriptions)



def makeProject(chromaQuery : str, chromaScenario : Collection) -> OneDesc :

    print(f"\nScenario query: ({chromaQuery})\n")
    queryResult = chromaScenario.query(query_texts=[chromaQuery], n_results=3)

    idx = -1
    numberChosen = 0
    distList = []
    combinedDoco = ""
    for distFloat in queryResult["distances"][0] :
        idx += 1
        if (distFloat > ConfigSingleton().conf["rag_scenario"]) : 
            break

        oneResult = OneQueryResult(
                id = queryResult["ids"][0][idx],
                name = queryResult["metadatas"][0][idx]["docName"], 
                desc = queryResult["documents"][0][idx], 
                query = chromaQuery,
                distance=distFloat 
            )

#            DebugUtils.dumpPydanticObject(oneResult, "one query result")
#            if DebugUtils.pressKey("Press c to move to next scenario:"):
#                break

        distList.append(distFloat)
        docText = ""
        if (queryResult["documents"]) :
            docText = queryResult["documents"][0][idx]
        metaInf = ""
        if (queryResult["metadatas"]) :
            metaInf = queryResult["metadatas"][0][idx]["docName"]
        combinedDoco = combinedDoco + "\n" + docText
        numberChosen = numberChosen + 1

    if not numberChosen :
        print(f"ERROR: cannot find ChromaDB records under distance {ConfigSingleton().conf["rag_scenario"]}")
        return None
    
#        print(f"Selected {numberChosen} scenarios from ChromaDB. Distances: {distList}")

    systemPrompt = f"""
        You are an expert technical writer. 
        Create title and description of the project based on information supplied. 
        Create a full paragraph for description.
        Do not format text. Remove line feeds and carriage returns.
        Output only the result.
        Here is the JSON schema for the OneDesc model you must use as context for what information is expected:
        {json.dumps(OneDesc.model_json_schema(), indent=2)}
        """
    prompt = f"{combinedDoco}"
    print(f"-------prompt---------\n{prompt}\n-------------\n")

    ollModel = OpenAIModel(model_name=ConfigSingleton().conf["main_llm_name"], 
                        provider=OpenAIProvider(base_url=ConfigSingleton().conf["llm_base_url"]))
    
    agent = Agent(ollModel,
                output_type=OneDesc,
                system_prompt = systemPrompt,
                retries=5,
                output_retries=5)
    try:
        result = agent.run_sync(prompt)
        oneDesc = OneDesc.model_validate_json(result.output.model_dump_json())
        oneDesc.description = oneDesc.description.replace("\n", "")
        DebugUtils.dumpPydanticObject(oneDesc, "Project Description")
    except pydantic_ai.exceptions.UnexpectedModelBehavior:
        print(f"Exception: pydantic_ai.exceptions.UnexpectedModelBehavior")
        return None
    return oneDesc


def extractInfoFromJobAd(jobInfo : OneRecord) -> OneResultList :

    systemPromptPhase1 = f"""
        You are an expert in cyber security, information technology and software development 
        You will be supplied text of job advertisement.
        Your job is to extract information from the text that matches user's request.
        """

    ollModel = OpenAIModel(model_name=ConfigSingleton().conf["main_llm_name"], 
                        provider=OpenAIProvider(base_url=ConfigSingleton().conf["llm_base_url"]))
    agent = Agent(ollModel, 
                system_prompt = systemPromptPhase1,
                retries=5,
                output_retries=5)

    promptPhase1 = f"""Extract the list of 
    activities, technologies, methodologies, software services, and software products from the text below.
    Combine all items in common list of strings.
    Do not separate by category.
    Avoid single word items.
    Output only lower-case characters.
    Output only the result.
    \n
    {jobInfo.description}
            """

    try:
        result = agent.run_sync(promptPhase1)
#        print(f"{result.output}")
    except pydantic_ai.exceptions.UnexpectedModelBehavior as e:
        print(f"extractInfo: Skipping due to exception: {e}")
        return None

    phase2Input = result.output

    systemPromptPhase2 = f"""
        You are an expert in JSON processing.
        Input prompt contains list in JSON format. 

        Here is the JSON schema for the OneResultList model you must use as context for what information is expected:
        {json.dumps(OneResultList.model_json_schema(), indent=2)}
        """
    promptPhase2 = f"""
                {phase2Input}
            """

    agentPhase2 = Agent(ollModel, 
                output_type=OneResultList,
                system_prompt = systemPromptPhase2,
                retries=5,
                output_retries=5)
    try:
        result = agentPhase2.run_sync(promptPhase2)
        oneResultList = OneResultList.model_validate_json(result.output.model_dump_json())
        DebugUtils.dumpPydanticObject(oneResultList, "list from job ad")
    except pydantic_ai.exceptions.UnexpectedModelBehavior as e:
#        print(f"extractInfoFromJobAd: Skipping due to exception: {e}")
        return None
    return oneResultList


def mapToActivity(oneResultList : OneResultList, chromaDBCollection : Collection) -> OneResultList :
    itemSet = set()
    for itm in oneResultList.results_list:
        listNew = getChromaDBMatchActivity(chromaDBCollection, itm)
        for itemNew in listNew.results_list:
            itemSet.add(itemNew)
    oneResultList = OneResultList(results_list = list(itemSet))
    DebugUtils.dumpPydanticObject(oneResultList, "Mapped list")
    return oneResultList

def mapToActivityFake(oneResultList : OneResultList, chromaDBCollection : Collection) -> OneResultList :
    oneResultList = OneResultList(results_list = [
            "Delivered Mitre ATT&CK framework implementation.",
            "Aligned security measures with business objectives.",
            "Led ISO 27001/02 compliance efforts.",
            "Applied best practices in cyber security to improve SOC capability.",
            "Contributed to strategic planning and cyber risk management at leadership level.",
            "Monitored and analysed security events.",
            "Performed threat hunting.",
            "Led incident response efforts.",
            "Utilized Splunk SIEM product.",
            "Performed tool optimization.",
            "Examined the effectiveness of established incident response policies and procedures.",
            "Developed threat models for multiple designs.",
            "Demonstrated SIEM platform proficiency.",
            "Assessed emerging technologies and frameworks to strengthen security controls and resilience.",
            "Implemented and enhanced cyber defences for critical systems.",
            "Helped with vulnerability management.",
            "Contributed to digital forensic investigations and reporting.",
            "Led, mentored, and supported cyber defence and response team."
            ])
    DebugUtils.dumpPydanticObject(oneResultList, "Mapped list")
    return oneResultList


#
# open chromaDB Collection
#
def getChromaCollection(chromaClient : chromadb.PersistentClient, collName : str) -> Collection :

    ef = OllamaEmbeddingFunction(
        model_name=ConfigSingleton().conf["rag_embed_llm"],
        url=ConfigSingleton().conf["rag_embed_url"]    
    )

    try:
        chromaColl = chromaClient.get_collection(
            name=collName,
            embedding_function=ef
        )
        print(f"Collection {collName} opened with {chromaColl.count()} documents")
    except Exception as e:
        print(f"Exception: {e}")
        return None
    return chromaColl



def getChromaDBMatchActivity(chromaDBCollection : Collection, queryString : str) -> OneResultList :

    totals = set()

    queryResult = chromaDBCollection.query(query_texts=[queryString], n_results=1)
    cutDist = ConfigSingleton().conf["rag_distmatch"]
    resultIdx = -1
    for distFloat in queryResult["distances"][0] :
        resultIdx += 1
        docText = ""
        if (queryResult["documents"]) :
            docText = queryResult["documents"][0][resultIdx]

        if (distFloat > cutDist) :
            break

        totals.add(docText)

    if not len(totals) :
        print(f"Query {queryString} did not get matches less than {ConfigSingleton().conf["rag_distmatch"]}")

#        with open("out.txt", "a") as fileOut:
#            fileOut.write(f'{{\n\t"id": "","name": "{queryString}",\n\t"description": "{queryString.capitalize()}."\n}},')
        return OneResultList(results_list = [])

    return OneResultList(results_list=list(totals))



def extractExecSection(jobInfo : OneRecord) -> OneDesc :

    systemPromptPhase1 = f"""
        You are an expert in cyber security, information technology and software development 
        You will be supplied text of job advertisement.
        Your job is to extract information from the text that matches user's request.
        Here is the JSON schema for the OneDesc model you must use as context for what information is expected:
        {json.dumps(OneDesc.model_json_schema(), indent=2)}
        """

    ollModel = OpenAIModel(model_name=ConfigSingleton().conf["main_llm_name"], 
                        provider=OpenAIProvider(base_url=ConfigSingleton().conf["llm_base_url"]))
    agent = Agent(ollModel, 
                output_type=OneDesc,                  
                system_prompt = systemPromptPhase1,
                retries=10,
                output_retries=10)

    promptPhase1 = f"""Extract the required role suitable for CV from the text below.
    Make required role the title.
    Fill description with generic description of the required role in past tense suitable for CV.
    Do not add formatting.
    Output only the result.
    \n
    {jobInfo.description}
            """

    try:
        result = agent.run_sync(promptPhase1)
        oneDesc = OneDesc.model_validate_json(result.output.model_dump_json())
        oneDesc.description = oneDesc.description.replace("\n", "")
        DebugUtils.dumpPydanticObject(oneDesc, "Executive Summary")
    except pydantic_ai.exceptions.UnexpectedModelBehavior as e:
        print(f"extractInfo: Skipping due to exception: {e}")
        return None
    return oneDesc


def makeWordDoc(allDesc : AllDesc) :

    allEmployers = AllEmployers(employer_list = [])
    allEmployers.employer_list.append(OneEmployer(
        name = "Darlowie Security Consulting",
        blurb = "Principal Consultant",
        start = datetime(2024, 11, 1),
        end = datetime.now(),
        numJobs = 3
    ))
    allEmployers.employer_list.append(OneEmployer(
        name = "NCC Group APAC",
        blurb = "Principal Consultant",
        start = datetime(2014, 12, 1),
        end = datetime(2024, 10, 30),
        numJobs = 15
    ))
    allEmployers.employer_list.append(OneEmployer(
        name = "NCC Group USA",
        blurb = "Senior Consultant",
        start = datetime(2013, 5, 15),
        end = datetime(2014, 10, 30),
        numJobs = 10
    ))

    # Create a new Document
    doc = Document()

    title = doc.add_heading(level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titleRun = title.add_run('Anton Pavlov')
    font = titleRun.font
    font.Name = 'Cambria'
    font.size = Pt(20)

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run('Melbourne, Victoria  Ph: 0428065430  Email:XXXX@YYYY.com')

    exec_paragraph = doc.add_paragraph()
    exec_paragraph.add_run(allDesc.exec_section.title).bold = True
    exec_paragraph.add_run('\n')
    exec_paragraph.add_run(allDesc.exec_section.description)

    expHeading = doc.add_heading(level=2)
    expHeading.add_run('Experience')

    idxCurrentJob = 0
    idxGlobalJob = 0
    for emp in allEmployers.employer_list:
        date_name_para = doc.add_paragraph()
        date_name_para.add_run(emp.start.strftime('%b %Y') + '-' +  emp.end.strftime('%b %Y'))
        date_name_para.add_run(' ')
        date_name_para.add_run(emp.name).bold = True
        blurb_para = doc.add_paragraph()
        blurb_para.add_run(emp.blurb)
        for oneDesc in allDesc.project_list[idxGlobalJob:] :
            idxCurrentJob += 1
            if (idxCurrentJob > emp.numJobs) :
                idxCurrentJob = 0
                break
            list_paragraph = doc.add_paragraph(style='List Bullet')
            list_paragraph.add_run(oneDesc.title).bold = True
            list_paragraph.add_run('\n')
            list_paragraph.add_run(oneDesc.description)
            list_paragraph.add_run('\n')
            idxGlobalJob += 1

    # Save the document
    doc.save(allDesc.ad_name + ".resume.docx")



if __name__ == "__main__":
    main()
